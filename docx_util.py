#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
DOCX生成工具类
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT, WD_SECTION

class Style(object):
    """
    样式类
    """

    STYLE_ALIGN_LEFT    = WD_ALIGN_PARAGRAPH.LEFT       # 左对齐
    STYLE_ALIGN_RIGHT   = WD_ALIGN_PARAGRAPH.RIGHT      # 右对齐
    STYLE_ALIGN_CENTER  = WD_ALIGN_PARAGRAPH.CENTER     # 居中
    STYLE_ALIGN_JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY    # 两端对齐

    def __init__(self, fontSize=14, fontName=u'宋体', bold=False, italic=False, underline=False, color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.LEFT):
        """
        构造函数

        @param fontSize：字号
        @param fontName：字体
        @param bold：是否粗体
        @param italic：是否斜体
        @param underline：是否下划线
        @param color：颜色
        @param align：对齐方式，具体值参照Style类中的定义
        """

        self.fontSize  = Pt(fontSize)
        self.fontName  = fontName
        self.bold      = bold
        self.italic    = italic
        self.color     = RGBColor(color[0], color[1], color[2])
        self.underline = underline
        self.align     = align

class Docx(object):
    """
    Docx生成类
    """

    PAGE_HORIZONTAL = WD_ORIENT.LANDSCAPE
    PAGE_VERTICAL = WD_ORIENT.PORTRAIT

    def  __init__(self):
        self._doc = Document()

    @staticmethod
    def CreateStyle(fontSize=14, fontName=u'宋体', bold=False, italic=False, underline=False, color=(0, 0 , 0), align=Style.STYLE_ALIGN_LEFT):
        """
        创建样式
        
        @param fontSize：字号
        @param fontName：字体
        @param bold：是否粗体
        @param italic：是否斜体
        @param underline：是否下划线
        @param color：颜色
        @param align：对齐方式，具体值参照Style类中的定义
        
        @return：创建的样式 
        """

        return Style(fontSize, fontName, bold, italic, underline, color, align)

    def AddHeader(self, level, text, style=None):
        """
        添加段落
        
        @param style：样式
        """

        self._doc.add_paragraph("", self._doc.styles['heading %d' % level])
        self.AddText(text, style)

    def AddParagraph(self, style=None):
        """
        添加段落
        
        @param style：样式
        """

        paragraph = self._doc.add_paragraph()
        if style:
            paragraph.paragraph_format.alignment = style.align

    def AddText(self, text='', style=None, pid=-1):
        """
        添加文字

        @param text：文字
        @param style：样式
        """

        p = self._doc.paragraphs[pid]
        run = p.add_run(text)
        if style:
            run.font.size = style.fontSize
            run.font.name = style.fontName
            run.font.color.rgb = style.color
            run.italic = style.italic
            run.underline = style.underline
            run.bold = style.bold
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), style.fontName)

    def AddPicture(self, picPath, width, dpi, caption, style):
        """
        添加图片

        @param picPath：图片路径
        @param width：图片宽度，设置了图片宽度后，高度根据图片大小自适应
        @param caption：题注
        @param style：样式
        """
        paragraph = self._doc.add_paragraph()
        paragraph.add_run().add_picture(picPath, width=Inches(float(width)/float(dpi)))
        paragraph.paragraph_format.alignment=style.align
        paragraph.paragraph_format.space_after = 0

        self.AddParagraph(style)
        self.AddText(caption, style)

    def AddTable(self, rows, cols, caption, style=None):
        """
        添加表格
        
        @param rows：行数
        @param cols：列数
        @param caption：题注
        @param style：样式
        """

        paragraph = self._doc.add_paragraph()
        if style:
            paragraph.paragraph_format.alignment=style.align
        paragraph.paragraph_format.space_after = 0
        self.AddText(caption, style)

        self._doc.add_table(rows=rows, cols=cols, style=self._doc.styles['Table Grid'])

    def SetCell(self, row, col, text, style=None, tid=-1):
        """
        设置单元格内容

        @param tId：表格ID
        @param row：行数
        @param col：列数
        @param text：文字
        @param style：样式
        """

        paragraph = self._doc.tables[tid].rows[row].cells[col].paragraphs[0]
        run=paragraph.add_run(text)
        if style:
            paragraph.paragraph_format.alignment = style.align
            run.font.size = style.fontSize
            run.font.name = style.fontName
            run.font.color.rgb = style.color
            run.italic = style.italic
            run.underline = style.underline
            run.bold = style.bold
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), style.fontName)

    def AddPageBreak(self):
        """
        添加分页符
        """

        self._doc.add_page_break()

    def AddSection(self, orientation):
        section = self._doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = orientation
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    def Save(self, docPath):
        """
        保存文档
        
        @param docPath：保存路径
        """

        self._doc.save(docPath)
        
if __name__ == "__main__":
    doc = Docx()
    style = doc.CreateStyle(fontSize=40, fontName=u'宋体', align=Style.STYLE_ALIGN_CENTER)
    pIndex = doc.AddParagraph(style)
    doc.AddText(u"交叉定标报告", style)

    style = doc.CreateStyle(fontName=u'宋体')
    doc.AddHeader(1, u"1.\t定标结果", style)

    doc.AddParagraph(style)
    style = doc.CreateStyle(fontName=u'宋体')
    doc.AddText(u"正常文字", style)
    style = doc.CreateStyle(fontName=u'宋体', bold=True)
    doc.AddText(u"粗体", style)
    style = doc.CreateStyle(fontName=u'宋体', italic=True)
    doc.AddText(u"斜体", style)
    style = doc.CreateStyle(fontName=u'宋体', underline=True)
    doc.AddText(u"下划线", style)
    style = doc.CreateStyle(fontName=u'宋体', bold=True, italic=True, underline=True)
    doc.AddText(u"粗体+斜体+下划线", style)
    style = doc.CreateStyle(fontName=u'宋体', color=(255, 0, 0))
    doc.AddText(u"红色", style)

    style = doc.CreateStyle(fontSize=12, fontName=u'宋体', italic=True, align=Style.STYLE_ALIGN_CENTER)
    doc.AddTable(7, 3, u"定标结果", style)

    # 添加表头
    style = doc.CreateStyle(fontName=u'宋体', bold=True, align=Style.STYLE_ALIGN_CENTER)
    doc.SetCell(0, 0, u"通道", style)
    doc.SetCell(0, 1, "BIAS", style)
    doc.SetCell(0, 2, "STD", style)

    style = doc.CreateStyle(fontName=u'宋体', align=Style.STYLE_ALIGN_CENTER)
    doc.SetCell(1, 0, "B6", style)
    doc.SetCell(2, 0, "B7", style)
    doc.SetCell(3, 0, "B8", style)
    doc.SetCell(4, 0, "B6", style)
    doc.SetCell(5, 0, "B7", style)
    doc.SetCell(6, 0, "B8", style)

    style = doc.CreateStyle(fontName=u'宋体', align=Style.STYLE_ALIGN_RIGHT)
    doc.SetCell(1, 1, "1.2K", style)
    doc.SetCell(2, 1, "1.3K", style)
    doc.SetCell(3, 1, "-1.0K", style)
    doc.SetCell(4, 1, "0.4K", style)
    doc.SetCell(5, 1, "0.8K", style)
    doc.SetCell(6, 1, "1.0K", style)

    doc.SetCell(1, 2, "1.2K", style)
    doc.SetCell(2, 2, "1.3K", style)
    doc.SetCell(3, 2, "-1.0K", style)
    doc.SetCell(4, 2, "0.4K", style)
    doc.SetCell(5, 2, "0.8K", style)
    doc.SetCell(6, 2, "1.0K", style)

    doc.AddSection()
    style = doc.CreateStyle(fontName=u'宋体')
    doc.AddHeader(1, u"2.\t结果分析", style)

    style = doc.CreateStyle(fontSize=12, fontName=u'宋体', italic=True, align=Style.STYLE_ALIGN_CENTER)
    doc.AddPicture("FY4A_IASIA_BIAS_20170701_20170731_B09_0625.png", 500, 100, u"B09通道偏差图", style)

    doc.AddPageBreak()

    style = doc.CreateStyle(fontSize=12, fontName=u'宋体', align=Style.STYLE_ALIGN_RIGHT)
    doc.AddParagraph(style)
    doc.AddText(u"北京华云星地通", style)

    style = doc.CreateStyle(fontSize=12, fontName=u'宋体', align=Style.STYLE_ALIGN_RIGHT)
    doc.AddParagraph(style)
    doc.AddText(u"2017-08-08", style)

    doc.Save('test1.docx')